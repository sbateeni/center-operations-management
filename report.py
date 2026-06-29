import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    # إنشاء مستند جديد
    doc = docx.Document()
    
    # ضبط هوامش الصفحة (1 إنش من كل اتجاه)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    doc.styles['Normal'].font.name = 'Traditional Arabic'
    doc.styles['Normal'].font.size = Pt(14)
    
    COLOR_PRIMARY = RGBColor(16, 44, 87)     # كحلي غامق وقور
    COLOR_SECONDARY = RGBColor(53, 114, 239)  # أزرق تقني
    COLOR_TEXT = RGBColor(40, 40, 40)        # رمادي داكن للنصوص لراحة العين
    COLOR_HIGHLIGHT = RGBColor(194, 137, 34)  # ذهبي دافئ للعناوين الجانبية والتنبيهات
    
    # ---------------------------------------------------------------------------
    # 1. تصميم صفحة الغلاف (Cover Page)
    # ---------------------------------------------------------------------------
    
    for _ in range(5):
        doc.add_paragraph()
        
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("____________________________________________________")
    run_line.font.color.rgb = COLOR_HIGHLIGHT
    run_line.font.bold = True
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("تقرير تفصيلي وشامل حول مشروع\nإدارة عمليات المراكز\n(Center Operations Management)")
    run_title.font.name = 'Traditional Arabic'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY
    
    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line2 = p_line2.add_run("____________________________________________________")
    run_line2.font.color.rgb = COLOR_HIGHLIGHT
    run_line2.font.bold = True
    
    for _ in range(3):
        doc.add_paragraph()
        
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run("إعداد: الضابط مقدم/ معتز عريدي\nنائب مدير مركز شرطة شقبا\nالتاريخ: يونيو 2026 م")
    run_meta.font.name = 'Traditional Arabic'
    run_meta.font.size = Pt(14)
    run_meta.font.bold = True
    run_meta.font.italic = True
    run_meta.font.color.rgb = COLOR_TEXT
    
    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # دوال مساعدة
    # ---------------------------------------------------------------------------
    
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Traditional Arabic'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(12)
        run_sub = p_sub.add_run("━" * 40)
        run_sub.font.color.rgb = COLOR_HIGHLIGHT
        run_sub.font.size = Pt(8)
        
    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Traditional Arabic'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        
    def add_body_paragraph(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.25
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            run_prefix = p.add_run(bold_prefix)
            run_prefix.font.name = 'Traditional Arabic'
            run_prefix.font.size = Pt(14)
            run_prefix.font.bold = True
            run_prefix.font.color.rgb = COLOR_PRIMARY
        run_text = p.add_run(text)
        run_text.font.name = 'Traditional Arabic'
        run_text.font.size = Pt(14)
        run_text.font.color.rgb = COLOR_TEXT
        return p

    def add_bullet_point(bold_title, description):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_title = p.add_run(bold_title + ": ")
        run_title.font.name = 'Traditional Arabic'
        run_title.font.size = Pt(13)
        run_title.font.bold = True
        run_title.font.color.rgb = COLOR_SECONDARY
        run_desc = p.add_run(description)
        run_desc.font.name = 'Traditional Arabic'
        run_desc.font.size = Pt(13)
        run_desc.font.color.rgb = COLOR_TEXT

    # ---------------------------------------------------------------------------
    # أولاً: مقدمة تمهيدية
    # ---------------------------------------------------------------------------
    add_heading_1("أولاً: مقدمة تمهيدية")
    add_body_paragraph("تعد إدارة العمليات الميدانية وتوجيه الحملات والوحدات بدقة وفي الوقت الفعلي من أبرز التحديات التي تواجه الأجهزة الأمنية والشرطية. في ظل تسارع الأحداث وتعدد البلاغات، لم يعد الاعتماد على الوسائل التقليدية في التوجيه اللاسلكي فقط كافياً للسيطرة التامة واستغلال الموارد المتاحة بالشكل الأمثل.")
    add_body_paragraph("استجابة لهذه المتطلبات، يأتي مشروع \"إدارة عمليات المراكز (Center Operations Management)\" ليوفر منصة قيادة وسيطرة متكاملة. يعتمد النظام على أحدث التقنيات الجغرافية وقواعد البيانات السحابية اللحظية (Supabase & React) لتمكين قادة العمليات من رؤية الوحدات وتتبع تحركاتها على الخرائط الرقمية بدقة عالية، مما يسهم في اتخاذ قرارات سريعة، مدروسة، ومبنية على بيانات حية من أرض الواقع.")

    # ---------------------------------------------------------------------------
    # ثانياً: الآلية المعمارية والتقنية
    # ---------------------------------------------------------------------------
    add_heading_1("ثانياً: الآلية المعمارية والتقنية للمشروع")
    add_body_paragraph("يعتمد النظام على بنية برمجية حديثة ومرنة تضمن السرعة والاستقرار. تتلخص آلية عمل النظام في المحاور التقنية التالية:")
    
    add_heading_2("1. الواجهة الجغرافية التفاعلية (Interactive Geographic Interface)")
    add_body_paragraph("يعتمد النظام على مكتبة (Leaflet) لبناء خرائط تفاعلية متقدمة تتيح للمستخدمين:")
    add_bullet_point("• العرض المكاني والطبقات", "دعم كامل لعرض الخرائط العادية وصور الأقمار الصناعية (Satellite View) لمناطق الاختصاص، مع إمكانية التبديل الفوري بينهما.")
    add_bullet_point("• تحديد النقاط الحيوية (Notes & Pins)", "إضافة علامات (Markers) توضح مواقع البلاغات أو الأهداف الميدانية بنقرة واحدة (Map Click) مع إمكانية تعديلها وحذفها والتحليق البصري نحوها (Fly-to-Target).")
    add_bullet_point("• التوجيه ورسم المسارات", "تحديد مسارات أولية وثانوية للوحدات الميدانية (Current & Secondary Routes) للوصول إلى الهدف بأقصر الطرق.")

    add_heading_2("2. التتبع الحي والمباشر (Real-time Tracking System)")
    add_body_paragraph("من خلال دمج قواعد بيانات (Supabase) ذات القدرة على بث التحديثات اللحظية (Real-time Subscriptions):")
    add_bullet_point("• التحديث الفوري للمواقع", "استقبال إحداثيات (GPS) الخاصة بالدوريات (Online Users) وعرضها على الخريطة دون الحاجة لإعادة تحميل الصفحة.")
    add_bullet_point("• حالة الوحدات الميدانية", "مراقبة حالة كل وحدة وتصنيفها لعرض جاهزيتها للتدخل (MyStatus)، مما يسهل على القيادة توجيه أقرب وحدة متاحة.")

    # ---------------------------------------------------------------------------
    # ثالثاً: المكونات المتقدمة لإدارة العمليات
    # ---------------------------------------------------------------------------
    add_heading_1("ثالثاً: المكونات المتقدمة لإدارة العمليات (Advanced Features)")
    add_body_paragraph("يحتوي الكود المصدري للنظام على وحدات (Components) متطورة لضمان إحكام السيطرة التامة وإدارة الأزمات باحترافية، ومن أبرزها:")
    
    add_heading_2("1. لوحة القيادة الاستراتيجية (Strategic Dashboard)")
    add_body_paragraph("تمنح القادة والمدراء واجهة موحدة توفر نظرة علوية شاملة (Dubai Strategic Dashboard Layer)، تتيح مراقبة كل الوحدات المتصلة، استعراض الملاحظات الأمنية المدونة في الميدان، وإدارة الحملات الجارية بكفاءة، ولا تظهر هذه اللوحة إلا لأصحاب الصلاحيات العليا.")

    add_heading_2("2. نظام الأدوار والصلاحيات (Role-Based Access Control)")
    add_body_paragraph("تم برمجة النظام للتعرف على تدرج الرتب والمسؤوليات:")
    add_bullet_point("• مستويات الوصول", "يدعم النظام صلاحيات متعددة (Super Admin, Governorate Admin, Center Admin, Admin) بحيث يرى كل مستوى ما يحتاج إليه فقط، من إنشاء المهام إلى تتبع الوحدات أو الاقتصار على تنفيذ المهام الميدانية.")

    add_heading_2("3. نظام الاستغاثة والطوارئ (SOS & Tactical Overlay)")
    add_body_paragraph("ميزة بالغة الأهمية لحماية العناصر الميدانية:")
    add_bullet_point("• إطلاق نداء استغاثة", "يتيح النظام للوحدة الميدانية إرسال إشارة (SOS) مشفرة، تظهر فوراً على شكل طبقة تكتيكية (Tactical Overlay) لدى غرف العمليات.")
    add_bullet_point("• تحديد المستغيث فوراً", "تتجه أنظار الخريطة والقيادة مباشرة إلى موقع المستغيث (Distressed User) لتوجيه الدعم اللوجستي أو العسكري فوراً.")

    add_heading_2("4. توجيه المهام والاعتراض (Dispatch & Intercept)")
    add_body_paragraph("توفير أدوات توجيه سريعة لقادة العمليات تتضمن:")
    add_bullet_point("• إرسال الأوامر الآلية", "تكليف الوحدات (Assignments) بالانتقال لموقع محدد (Dispatch Target Location) بضغطة زر وتلقي إشعارات بالقبول.")
    add_bullet_point("• ميزة الاعتراض والملاحقة", "إمكانية إعطاء أمر مباشر لوحدة ميدانية لاعتراض (Intercept) هدف معين يتحرك على الخريطة ضمن حملة أمنية منظمة (Active Campaign).")

    # ---------------------------------------------------------------------------
    # رابعاً: المزايا الاستراتيجية والتشغيلية
    # ---------------------------------------------------------------------------
    add_heading_1("رابعاً: المزايا الاستراتيجية والتشغيلية")
    add_body_paragraph("يوفر هذا النظام نقلة نوعية في كفاءة الاستجابة الأمنية عبر المزايا التالية:")
    add_bullet_point("1. سرعة الاستجابة ورفع الكفاءة", "توجيه أقرب وحدة متاحة لموقع الحدث يقلل من زمن الاستجابة بشكل ملحوظ.")
    add_bullet_point("2. الرؤية البانورامية الشاملة", "منح القيادة صورة كاملة وموحدة لمسرح العمليات مما يسهل إدارة الأزمات وتوزيع القوات.")
    add_bullet_point("3. التوثيق والمحاسبية", "أرشفة وتوثيق جميع المهام (Logs) بشكل يضمن القدرة على العودة لتفاصيل أي حملة أمنية لأغراض التدريب والتقييم.")
    add_bullet_point("4. مرونة التشغيل والتصميم المتجاوب", "بناء الواجهات لتعمل على الشاشات العملاقة في غرف العمليات وعلى الهواتف المحمولة في الميدان بنفس الكفاءة وسرعة الأداء.")

    # ---------------------------------------------------------------------------
    # خامساً: التحديات التقنية والتشغيلية
    # ---------------------------------------------------------------------------
    add_heading_1("خامساً: التحديات التقنية والتشغيلية")
    add_body_paragraph("لضمان استمرارية وكفاءة النظام في أصعب الظروف الميدانية، يجب التعامل بحزم مع التحديات التالية:")
    
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    shading_elm = parse_xml(r'<w:shd {} w:fill="F4F4F6"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'''
        <w:tcBorders {}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="single" w:sz="36" w:space="0" w:color="C28922"/>
        </w:tcBorders>
    '''.format(nsdecls('w')))
    tcPr.append(tcBorders)
    
    p_box = cell.paragraphs[0]
    p_box.paragraph_format.space_before = Pt(6)
    p_box.paragraph_format.space_after = Pt(6)
    
    r_box_title = p_box.add_run("⚠️ تحديات استراتيجية هامة للتشغيل:\n\n")
    r_box_title.font.name = 'Traditional Arabic'
    r_box_title.font.size = Pt(13)
    r_box_title.font.bold = True
    r_box_title.font.color.rgb = COLOR_HIGHLIGHT
    
    r_box_desc = p_box.add_run(
        "1. استقرار شبكات الاتصال الميدانية: يعتمد النظام اللحظي على توفر اتصال إنترنت موثوق لدى الوحدات في الميدان (3G/4G/5G) لضمان دقة بث الإحداثيات (GPS).\n"
        "2. أمن البيانات السحابية والمصادقة: يتطلب تأمين واجهات الدخول (AuthPage) وعمليات الاعتماد (Pending Approval) لحماية لوحة القيادة من الوصول غير المصرح به.\n"
        "3. استهلاك البطارية للأجهزة الميدانية: التحديث اللحظي لنقاط الخريطة يستهلك الموارد، مما يوجب توفير مصادر طاقة مستدامة للأجهزة المحمولة داخل الدوريات."
    )
    r_box_desc.font.name = 'Traditional Arabic'
    r_box_desc.font.size = Pt(12)
    r_box_desc.font.color.rgb = COLOR_TEXT
    
    doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # سادساً: الخلاصة والتوصيات
    # ---------------------------------------------------------------------------
    add_heading_1("سادساً: الخلاصة والتوصيات الختامية")
    add_body_paragraph("يمثل تطبيق \"إدارة عمليات المراكز\" المستند على تقنيات حديثة خطوة جوهرية نحو التحول الرقمي وإدارة الموارد الشرطية بشكل احترافي وذكي. فهو يتجاوز كونه أداة تتبع ليصبح عقلاً إلكترونياً مساعداً يجمع أطراف مسرح العمليات أمام صانع القرار في شاشة واحدة تتفاعل لحظياً مع التطورات.")
    add_body_paragraph("ولضمان نجاح تشغيل هذا النظام واستدامته، نوصي بما يلي:")
    add_bullet_point("• تأمين البنية التحتية للاتصالات", "توفير حلول اتصال بديلة أو شرائح اتصال آمنة ومستقرة للوحدات الميدانية.")
    add_bullet_point("• تدريب الأطقم والمشغلين", "تدريب شامل للمشغلين على كيفية استخدام لوحة القيادة الاستراتيجية وإدارة الأزمات ونداءات الاستغاثة عبر النظام.")
    add_bullet_point("• النسخ الاحتياطي والأمن السيبراني", "تنفيذ بروتوكولات حماية عالية للأنظمة لضمان عدم اختراق أو تسريب معلومات الحملات الجارية وحفظ السجلات للاستخدام المستقبلي.")

    # تطبيق خاصية RTL
    for p in doc.paragraphs:
        pPr = p._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
        for run in p.runs:
            rPr = run._element.get_or_add_rPr()
            rtl = OxmlElement('w:rtl')
            rtl.set(qn('w:val'), '1')
            rPr.append(rtl)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pPr = p._element.get_or_add_pPr()
                    bidi = OxmlElement('w:bidi')
                    bidi.set(qn('w:val'), '1')
                    pPr.append(bidi)
                    for run in p.runs:
                        rPr = run._element.get_or_add_rPr()
                        rtl = OxmlElement('w:rtl')
                        rtl.set(qn('w:val'), '1')
                        rPr.append(rtl)

    return doc

# تشغيل الدالة لإنشاء المستند
doc = create_document()
doc.save("Operations_Control_Room_Report.docx")
